"""
Document Ingestion Pipeline for Copilot RSSI
Processes PDF, Word, Excel files and prepares them for vector storage
"""
import fitz  # PyMuPDF
import docx
import openpyxl
import pytesseract
from PIL import Image
import io
import re
from typing import List, Dict, Tuple
import os

class DocumentIngestion:
    def __init__(self):
        """Initialize document ingestion pipeline"""
        self.chunk_size = 1000  # characters per chunk
        self.chunk_overlap = 200  # characters overlap between chunks
    
    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        ext = filename.split(".")[-1].lower()
        text = ""
        
        try:
            if ext == "pdf":
                text = self._extract_pdf(file_bytes)
            elif ext in ["docx", "doc"]:
                text = self._extract_docx(file_bytes)
            elif ext in ["xlsx", "xls"]:
                text = self._extract_excel(file_bytes)
            elif ext in ["png", "jpg", "jpeg", "tiff", "bmp"]:
                text = self._extract_image(file_bytes)
            elif ext == "txt":
                text = file_bytes.decode('utf-8', errors='ignore')
            else:
                print(f"Unsupported file format: {ext}")
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
        
        return text
    
    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF with OCR fallback"""
        text = ""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                page_text = page.get_text()
                text += page_text + "\n"
            
            # OCR fallback if no text extracted
            if not text.strip():
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    text += pytesseract.image_to_string(Image.open(io.BytesIO(img_data))) + "\n"
        except Exception as e:
            print(f"PDF extraction error: {e}")
        
        return text
    
    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from Word documents"""
        text = ""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
        except Exception as e:
            print(f"DOCX extraction error: {e}")
        
        return text
    
    def _extract_excel(self, file_bytes: bytes) -> str:
        """Extract text from Excel files"""
        text = ""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    text += " | ".join(row_text) + "\n"
        except Exception as e:
            print(f"Excel extraction error: {e}")
        
        return text
    
    def _extract_image(self, file_bytes: bytes) -> str:
        """Extract text from images using OCR"""
        text = ""
        try:
            text = pytesseract.image_to_string(Image.open(io.BytesIO(file_bytes)))
        except Exception as e:
            print(f"Image OCR error: {e}")
        
        return text
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict[str, str]]:
        """
        Split text into chunks for vector storage
        Returns list of dicts with keys: id, text, metadata
        """
        if metadata is None:
            metadata = {}
        
        chunks = []
        text_length = len(text)
        
        for i in range(0, text_length, self.chunk_size - self.chunk_overlap):
            chunk_text = text[i:i + self.chunk_size]
            
            # Skip empty chunks
            if not chunk_text.strip():
                continue
            
            chunk_id = f"chunk_{i}_{metadata.get('filename', 'unknown')}"
            
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_index"] = len(chunks)
            chunk_metadata["chunk_start"] = i
            chunk_metadata["chunk_end"] = min(i + self.chunk_size, text_length)
            
            chunks.append({
                "id": chunk_id,
                "text": chunk_text,
                "metadata": chunk_metadata
            })
        
        return chunks
    
    def process_file(self, file_path: str, filename: str = None) -> List[Dict[str, str]]:
        """
        Process a single file and return chunks
        """
        if filename is None:
            filename = os.path.basename(file_path)
        
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
        
        text = self.extract_text(file_bytes, filename)
        
        metadata = {
            "filename": filename,
            "file_path": file_path,
            "file_type": filename.split(".")[-1].lower(),
            "total_chars": len(text)
        }
        
        return self.chunk_text(text, metadata)
    
    def process_directory(self, directory: str) -> List[Dict[str, str]]:
        """
        Process all supported files in a directory
        """
        all_chunks = []
        supported_extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.png', '.jpg', '.jpeg'}
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_extensions:
                    file_path = os.path.join(root, file)
                    try:
                        chunks = self.process_file(file_path, file)
                        all_chunks.extend(chunks)
                        print(f"Processed: {file} ({len(chunks)} chunks)")
                    except Exception as e:
                        print(f"Error processing {file}: {e}")
        
        return	all_chunks
    
    def process_referentiels(self, referentiels_dir: str) -> List[Dict[str, str]]:
        """
        Process referentiels (ISO/NIST/CIS/Loi) documents
        """
        all_chunks = []
        
        if not os.path.exists(referentiels_dir):
            print(f"Referentiels directory not found: {referentiels_dir}")
            return all_chunks
        
        for root, dirs, files in os.walk(referentiels_dir):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in ['.pdf', '.docx', '.doc', '.txt']:
                    file_path = os.path.join(root, file)
                    try:
                        chunks = self.process_file(file_path, file)
                        
                        # Add referentiel-specific metadata
                        for chunk in chunks:
                            chunk["metadata"]["type"] = "REFERENTIEL"
                            chunk["metadata"]["category"] = self._detect_referentiel_category(file)
                        
                        all_chunks.extend(chunks)
                        print(f"Processed referentiel: {file} ({len(chunks)} chunks)")
                    except Exception as e:
                        print(f"Error processing referentiel {file}: {e}")
        
        return all_chunks
    
    def _detect_referentiel_category(self, filename: str) -> str:
        """Detect the category of referentiel from filename"""
        filename_lower = filename.lower()
        
        if 'iso' in filename_lower:
            return 'ISO'
        elif 'nist' in filename_lower:
            return 'NIST'
        elif 'cis' in filename_lower:
            return 'CIS'
        elif 'loi' in filename_lower or '09-08' in filename_lower:
            return 'LOI'
        elif 'politique' in filename_lower or 'ssi' in filename_lower:
            return 'POLITIQUE'
        else:
            return 'AUTRE'
